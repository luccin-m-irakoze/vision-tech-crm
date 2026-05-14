"""
Build a single Word (.docx) submission from Phase 1–4 markdown sources.

Run from repo root:
  py -3 docs/build_submission_docx.py

Output:
  docs/SUBMISSION_Vision_Tech_CRM_All_Phases.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

DOCS_DIR = Path(__file__).resolve().parent
OUT_PATH = DOCS_DIR / "SUBMISSION_Vision_Tech_CRM_All_Phases.docx"


def set_run_font(run, name: str = "Calibri", size_pt: float | None = 11) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def add_formatted_paragraph(doc: Document, text: str) -> None:
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
            set_run_font(run)
        elif part:
            run = p.add_run(part)
            set_run_font(run)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size_pt=10)


def process_markdown_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                flush_code()
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|") and stripped.count("|") >= 2:
            rows: list[list[str]] = []
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line.startswith("|"):
                    break
                if re.match(r"^\|[\s\-:|]+\|\s*$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                rows.append(cells)
                i += 1
            add_markdown_table(doc, rows)
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:].strip())
            p.paragraph_format.left_indent = Pt(18)
            for run in p.runs:
                set_run_font(run, size_pt=10)
                run.italic = True
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_run_font(run)
        elif re.match(r"^\d+\.\s+", stripped):
            body = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(body, style="List Number")
            for run in p.runs:
                set_run_font(run)
        elif stripped:
            add_formatted_paragraph(doc, line)
        i += 1


def read_utf8(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = title.add_run(
        "Vision Technologies CRM\n"
        "Complete Submission — Phases 1 through 4\n"
        "SENG 8240 — Best Programming Practices and Design Patterns"
    )
    r.bold = True
    set_run_font(r, size_pt=16)

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = sub.add_run(
        "\nCase study: Vision Technologies Company Ltd (VTC)\n"
        "Student: MUNEZERO IRAKOZE Luccin · ID 25815\n"
        "Instructor: Rutarindwa Jean Pierre\n"
        "Academic year: 2025 / 2026 — Semester II"
    )
    set_run_font(r2, size_pt=12)

    note = doc.add_paragraph()
    rn = note.add_run(
        "\nNote on diagrams: Mermaid diagram sources appear as monospace blocks. "
        "For printed submission, render them at https://mermaid.live and paste images "
        "where needed."
    )
    set_run_font(rn, size_pt=10)
    rn.italic = True

    doc.add_page_break()


def main() -> int:
    phase2_intro = (
        "## Phase 2 — Answering the assignment questions\n\n"
        "This phase requires a summarized software development prototype following best "
        "practices, and a clearly identified design pattern. The prototype is the "
        "full-stack CRM described in the repository `README.md` (Vue 3 + Spring Boot 3 + "
        "MySQL 8 + Docker). The sections below document **coding standards** (Google "
        "Java Style Guide and Vue.js Style Guide), the **MVC** pattern as applied across "
        "backend and frontend, and the **folder structure** of the implementation.\n\n"
        "---\n\n"
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc)

    process_markdown_body(doc, read_utf8(DOCS_DIR / "PHASE1.md"))

    doc.add_page_break()
    doc.add_heading(
        "PHASE 2 — Software Prototype, Coding Standards & Design Pattern",
        level=0,
    )
    process_markdown_body(doc, phase2_intro)
    process_markdown_body(doc, read_utf8(DOCS_DIR / "CODING_STANDARDS.md"))
    process_markdown_body(doc, read_utf8(DOCS_DIR / "DESIGN_PATTERN.md"))
    process_markdown_body(doc, read_utf8(DOCS_DIR / "FOLDER_TREE.md"))

    doc.add_page_break()
    process_markdown_body(doc, read_utf8(DOCS_DIR / "GIT_GUIDE.md"))

    doc.add_page_break()
    process_markdown_body(doc, read_utf8(DOCS_DIR / "TEST_PLAN.md"))

    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
